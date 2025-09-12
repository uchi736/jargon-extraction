#!/usr/bin/env python3
"""
文書前処理スクリプト
各種文書を読み込んでテキストを抽出し、JSON形式で保存
"""

import json
import sys
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import hashlib

import pymupdf  # PyMuPDF
import docx
from bs4 import BeautifulSoup
import markdown
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table

console = Console()


class DocumentPreprocessor:
    """文書前処理クラス"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF',
        '.docx': 'Word',
        '.doc': 'Word',
        '.txt': 'Text',
        '.md': 'Markdown',
        '.html': 'HTML',
        '.htm': 'HTML'
    }
    
    def __init__(self, output_dir: Path = Path("preprocessed")):
        """
        初期化
        
        Args:
            output_dir: 前処理済みファイルの出力ディレクトリ
        """
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 処理統計
        self.stats = {
            "processed": 0,
            "failed": 0,
            "skipped": 0,
            "total_chars": 0
        }
    
    def preprocess_file(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """
        単一ファイルを前処理
        
        Args:
            file_path: 処理するファイルパス
            
        Returns:
            前処理結果の辞書、失敗時はNone
        """
        if not file_path.exists():
            console.print(f"[red]Error: File not found: {file_path}[/red]")
            return None
        
        suffix = file_path.suffix.lower()
        
        if suffix not in self.SUPPORTED_EXTENSIONS:
            console.print(f"[yellow]Skip: Unsupported format {suffix}: {file_path.name}[/yellow]")
            self.stats["skipped"] += 1
            return None
        
        try:
            # テキスト抽出
            if suffix == '.pdf':
                text = self._extract_pdf(file_path)
            elif suffix in ['.docx', '.doc']:
                text = self._extract_word(file_path)
            elif suffix == '.txt':
                text = self._extract_text(file_path)
            elif suffix == '.md':
                text = self._extract_markdown(file_path)
            elif suffix in ['.html', '.htm']:
                text = self._extract_html(file_path)
            else:
                text = None
            
            if text:
                # メタデータ生成
                stat = file_path.stat()
                content_hash = hashlib.sha256(text.encode()).hexdigest()[:16]
                
                result = {
                    "file_name": file_path.name,
                    "file_path": str(file_path.absolute()),
                    "file_type": self.SUPPORTED_EXTENSIONS[suffix],
                    "file_size": stat.st_size,
                    "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "processed_time": datetime.now().isoformat(),
                    "content_hash": content_hash,
                    "text_length": len(text),
                    "text": text
                }
                
                self.stats["processed"] += 1
                self.stats["total_chars"] += len(text)
                
                console.print(f"[green]OK[/green] {file_path.name} ({len(text):,} chars)")
                return result
            else:
                self.stats["failed"] += 1
                console.print(f"[red]NG[/red] {file_path.name} (Failed to extract text)")
                return None
                
        except Exception as e:
            self.stats["failed"] += 1
            console.print(f"[red]Error: {file_path.name}: {e}[/red]")
            return None
    
    def _extract_pdf(self, file_path: Path) -> str:
        """PDFからテキスト抽出"""
        doc = pymupdf.open(str(file_path))
        text_parts = []
        
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"[Page {page_num}]\n{text}")
        
        doc.close()
        
        if not text_parts:
            raise ValueError("PDFからテキストを抽出できませんでした")
        
        return "\n\n".join(text_parts)
    
    def _extract_word(self, file_path: Path) -> str:
        """Wordからテキスト抽出"""
        doc = docx.Document(str(file_path))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # テーブルのテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        if not paragraphs:
            raise ValueError("Word文書からテキストを抽出できませんでした")
        
        return "\n\n".join(paragraphs)
    
    def _extract_text(self, file_path: Path) -> str:
        """テキストファイル読み込み（エンコーディング自動検出）"""
        encodings = ['utf-8', 'shift-jis', 'euc-jp', 'iso-2022-jp', 'cp932']
        
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        
        # 全て失敗した場合はエラー無視で読み込み
        return file_path.read_text(encoding='utf-8', errors='ignore')
    
    def _extract_markdown(self, file_path: Path) -> str:
        """Markdownからプレーンテキスト抽出"""
        content = self._extract_text(file_path)
        
        # MarkdownをHTMLに変換してからテキスト抽出
        html = markdown.markdown(content, extensions=['tables', 'fenced_code', 'footnotes'])
        soup = BeautifulSoup(html, 'html.parser')
        
        # コードブロックは保持
        for code in soup.find_all('code'):
            code.string = f"[CODE: {code.get_text()}]"
        
        return soup.get_text()
    
    def _extract_html(self, file_path: Path) -> str:
        """HTMLからテキスト抽出"""
        content = self._extract_text(file_path)
        soup = BeautifulSoup(content, 'html.parser')
        
        # スクリプトとスタイルタグを除去
        for element in soup(['script', 'style', 'meta', 'link']):
            element.decompose()
        
        # テキストを抽出
        text = soup.get_text()
        
        # 空行を整理
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line)
    
    def preprocess_directory(
        self, 
        input_dir: Path, 
        extensions: Optional[List[str]] = None,
        recursive: bool = False
    ) -> Path:
        """
        ディレクトリ内のファイルを一括前処理
        
        Args:
            input_dir: 入力ディレクトリ
            extensions: 処理する拡張子リスト
            recursive: サブディレクトリも処理するか
            
        Returns:
            出力ファイルのパス
        """
        if not input_dir.is_dir():
            console.print(f"[red]Error: Directory not found: {input_dir}[/red]")
            return None
        
        # 拡張子リストの設定
        if extensions is None:
            extensions = list(self.SUPPORTED_EXTENSIONS.keys())
        
        # ファイルを検索
        files = []
        for ext in extensions:
            if recursive:
                pattern = f"**/*{ext}"
            else:
                pattern = f"*{ext}"
            
            files.extend(input_dir.glob(pattern))
            files.extend(input_dir.glob(pattern.upper()))
        
        # 重複を除去
        files = list(set(files))
        
        if not files:
            console.print(f"[yellow]Warning: No processable files in {input_dir}[/yellow]")
            return None
        
        console.print(f"[cyan]Processing: {len(files)} files[/cyan]")
        
        # 処理結果を格納
        results = []
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console
        ) as progress:
            task = progress.add_task(
                "[cyan]Preprocessing...",
                total=len(files)
            )
            
            for file_path in sorted(files):
                result = self.preprocess_file(file_path)
                if result:
                    results.append(result)
                progress.update(task, advance=1)
        
        # 結果を保存
        if results:
            output_file = self._save_results(results, input_dir)
            self._display_statistics()
            return output_file
        else:
            console.print("[red]No processable files found[/red]")
            return None
    
    def _save_results(self, results: List[Dict], source_dir: Path) -> Path:
        """
        処理結果をJSON形式で保存
        
        Args:
            results: 処理結果リスト
            source_dir: 元のディレクトリ
            
        Returns:
            出力ファイルパス
        """
        # 出力ファイル名を生成
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        source_name = source_dir.name if source_dir.name != "." else "current"
        output_file = self.output_dir / f"{source_name}_{timestamp}.json"
        
        # メタデータを追加
        output_data = {
            "metadata": {
                "source_directory": str(source_dir.absolute()),
                "processed_at": datetime.now().isoformat(),
                "total_files": len(results),
                "total_characters": sum(r["text_length"] for r in results),
                "preprocessor_version": "1.0.0"
            },
            "documents": results
        }
        
        # JSON形式で保存
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        
        console.print(f"\n[green]Saved: {output_file}[/green]")
        
        # インデックスファイルも更新
        self._update_index(output_file, output_data["metadata"])
        
        return output_file
    
    def _update_index(self, output_file: Path, metadata: Dict):
        """
        インデックスファイルを更新
        
        Args:
            output_file: 出力ファイル
            metadata: メタデータ
        """
        index_file = self.output_dir / "index.json"
        
        # 既存のインデックスを読み込み
        if index_file.exists():
            with open(index_file, 'r', encoding='utf-8') as f:
                index = json.load(f)
        else:
            index = {"preprocessed_files": []}
        
        # 新しいエントリを追加
        entry = {
            "file": output_file.name,
            "source": metadata["source_directory"],
            "processed_at": metadata["processed_at"],
            "files": metadata["total_files"],
            "characters": metadata["total_characters"]
        }
        
        index["preprocessed_files"].append(entry)
        
        # 保存
        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(index, f, ensure_ascii=False, indent=2)
    
    def _display_statistics(self):
        """処理統計を表示"""
        table = Table(title="Preprocessing Statistics")
        table.add_column("Item", style="cyan")
        table.add_column("Value", justify="right")
        
        table.add_row("Processed", f"[green]{self.stats['processed']}[/green]")
        table.add_row("Failed", f"[red]{self.stats['failed']}[/red]")
        table.add_row("Skipped", f"[yellow]{self.stats['skipped']}[/yellow]")
        table.add_row("Total Chars", f"{self.stats['total_chars']:,}")
        
        console.print(table)
    
    @staticmethod
    def list_preprocessed(output_dir: Path = Path("preprocessed")):
        """
        前処理済みファイルの一覧を表示
        
        Args:
            output_dir: 前処理済みファイルのディレクトリ
        """
        index_file = output_dir / "index.json"
        
        if not index_file.exists():
            console.print("[yellow]No preprocessed files found[/yellow]")
            return
        
        with open(index_file, 'r', encoding='utf-8') as f:
            index = json.load(f)
        
        table = Table(title="Preprocessed Files")
        table.add_column("File", style="cyan")
        table.add_column("Source", style="green")
        table.add_column("Date")
        table.add_column("Files", justify="right")
        table.add_column("Chars", justify="right")
        
        for entry in index["preprocessed_files"]:
            source_name = Path(entry["source"]).name
            date_str = entry["processed_at"].split("T")[0]
            table.add_row(
                entry["file"],
                source_name,
                date_str,
                str(entry["files"]),
                f"{entry['characters']:,}"
            )
        
        console.print(table)


def main():
    """メイン処理"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Document Preprocessing Tool")
    parser.add_argument("input", type=Path, nargs='?', help="Input directory or file")
    parser.add_argument("-o", "--output", type=Path, default=Path("preprocessed"),
                       help="Output directory (default: preprocessed)")
    parser.add_argument("-r", "--recursive", action="store_true",
                       help="Process subdirectories recursively")
    parser.add_argument("-l", "--list", action="store_true",
                       help="List preprocessed files")
    parser.add_argument("-e", "--extensions", nargs="+",
                       help="File extensions to process (e.g., .pdf .docx)")
    
    args = parser.parse_args()
    
    if args.list:
        DocumentPreprocessor.list_preprocessed(args.output)
        return
    
    if not args.input:
        parser.error("Input path is required unless using --list")
    
    # 前処理実行
    preprocessor = DocumentPreprocessor(output_dir=args.output)
    
    if args.input.is_file():
        # 単一ファイル処理
        result = preprocessor.preprocess_file(args.input)
        if result:
            output_file = preprocessor._save_results([result], args.input.parent)
            console.print(f"[green]Complete: {output_file}[/green]")
    else:
        # ディレクトリ処理
        preprocessor.preprocess_directory(
            args.input,
            extensions=args.extensions,
            recursive=args.recursive
        )


if __name__ == "__main__":
    main()