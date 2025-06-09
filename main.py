#!/usr/bin/env python3
"""
専門用語・類義語辞書生成プログラム (2パスLLM・階層化対応版)
技術文書から専門用語を抽出し(パス1)、LLMによる全体統合を経て(パス2)、高精度な辞書を生成する
"""

import os
import sys
import json
import asyncio
import logging
from pathlib import Path
from typing import List, Dict, Optional, Set, Any, TypeVar, Callable
from dataclasses import dataclass, field
from datetime import datetime
from collections import defaultdict, Counter
import re
import unicodedata
import random

# 外部ライブラリ
import yaml
import typer
from rich.console import Console
from rich.table import Table
from dotenv import load_dotenv
import aiofiles

# ドキュメント処理
from unstructured.partition.auto import partition
from pypdf import PdfReader
import docx
from bs4 import BeautifulSoup
import markdown

# 自然言語処理
import numpy as np

# LLM関連 (Gemini)
import google.generativeai as genai
from google.generativeai.types import GenerationConfig, HarmCategory, HarmBlockThreshold
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from pydantic import BaseModel, Field as PydanticField

# ベクトルDB
import psycopg2
from pgvector.psycopg2 import register_vector

# API
from fastapi import FastAPI, HTTPException, Query, Depends
import uvicorn

# 環境変数読み込み
load_dotenv()

# --- ログ設定 ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('term_extraction.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# --- グローバル設定 ---
console = Console()
app = typer.Typer(help="専門用語・類義語辞書生成プログラム (2パスLLM・階層化対応版)")

# ==============================================================================
# SECTION 1: データモデル (Data Models)
# ==============================================================================

class ExtractedTerm(BaseModel):
    """LLMからの出力形式を定義するPydanticモデル"""
    headword: str = PydanticField(..., description="主要な専門用語（見出し語）")
    synonyms: List[str] = PydanticField(default_factory=list, description="類義語や表記揺れのリスト")
    definition: str = PydanticField(..., description="文脈に基づいた専門用語の簡潔な定義")
    category: Optional[str] = PydanticField(None, description="その用語が属する上位の概念やカテゴリ")

class TermList(BaseModel):
    """LLMが出力する用語リストのコンテナ"""
    terms: List[ExtractedTerm]

@dataclass
class Term:
    """アプリケーション内で使用する専門用語エンティティ"""
    headword: str
    synonyms: List[str] = field(default_factory=list)
    definition: str = ""
    category: Optional[str] = None
    confidence: float = 1.0
    source_docs: Set[str] = field(default_factory=set)
    embedding: Optional[np.ndarray] = field(default=None, repr=False)
    created_at: datetime = field(default_factory=datetime.now)

    def to_dict(self) -> Dict:
        """辞書形式に変換"""
        return {
            "headword": self.headword,
            "synonyms": self.synonyms,
            "definition": self.definition,
            "category": self.category,
            "confidence": self.confidence,
            "source_docs": list(self.source_docs),
            "created_at": self.created_at.isoformat()
        }

@dataclass
class Document:
    """処理対象文書"""
    path: Path
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[str] = field(default_factory=list)

# ==============================================================================
# SECTION 2: 設定管理 (Configuration)
# ==============================================================================

class Config:
    """設定管理クラス"""
    def __init__(self, config_path: str = "config.yml"):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                self.data = yaml.safe_load(f)
        except FileNotFoundError:
            logger.error(f"設定ファイルが見つかりません: {config_path}")
            console.print(f"[bold red]エラー: 設定ファイル '{config_path}' が見つかりません。[/bold red]")
            console.print("'python main.py init-project' を実行して初期化してください。")
            raise typer.Exit(1)

        # LLM設定
        self.llm_model = self.data.get('llm', {}).get('model', 'gemini-2.0-flash')
        self.embedding_model = self.data.get('embedding', {}).get('model', 'models/text-embedding-004')
        self.embedding_dim = self.data.get('embedding', {}).get('dim', 768)
        self.batch_size = self.data.get('embedding', {}).get('batch_size', 100)

        # Storage設定
        self.storage_driver = self.data.get('storage', {}).get('driver', 'postgres')
        self.dsn = os.getenv("DATABASE_URL", self.data.get('storage', {}).get('dsn', 'postgresql://user:password@localhost:5432/term_db'))
        self.use_pgvector = self.data.get('storage', {}).get('use_pgvector', True)

        # 並列処理
        self.parallelism = self.data.get('parallelism', 4)

        # APIキー設定
        self.google_api_key = os.getenv("GOOGLE_API_KEY")
        if not self.google_api_key:
            logger.error("GOOGLE_API_KEY環境変数が設定されていません")
            console.print("[bold red]エラー: GOOGLE_API_KEYが設定されていません。[/bold red]")
            console.print("\n以下の手順でAPIキーを設定してください：")
            console.print("1. Google AI Studio (https://aistudio.google.com/apikey) でAPIキーを取得")
            console.print("2. .envファイルに以下の形式で記載:")
            console.print("   GOOGLE_API_KEY=your-actual-api-key-here")
            raise typer.Exit(1)

        # Gemini APIの初期化
        try:
            genai.configure(api_key=self.google_api_key)
            logger.info("Gemini API設定完了")
        except Exception as e:
            logger.error(f"Gemini API設定エラー: {e}")
            raise

# ==============================================================================
# SECTION 3: ユーティリティ関数
# ==============================================================================

T = TypeVar('T')

async def retry_async(
    func: Callable[..., T],
    max_attempts: int = 3,
    initial_delay: float = 1.0,
    backoff_factor: float = 2.0,
    max_delay: float = 60.0,
    *args,
    **kwargs
) -> Optional[T]:
    """非同期関数のリトライラッパー（指数バックオフとジッター付き）"""
    last_exception = None
    for attempt in range(max_attempts):
        try:
            return await func(*args, **kwargs)
        except Exception as e:
            last_exception = e
            if attempt == max_attempts - 1:
                logger.error(f"リトライ失敗 ({attempt + 1}/{max_attempts}): {e}")
                raise
            delay = min(initial_delay * (backoff_factor ** attempt), max_delay)
            jitter = random.uniform(0, delay * 0.1)
            actual_delay = delay + jitter
            logger.warning(f"リトライ {attempt + 1}/{max_attempts}: {e}")
            logger.info(f"{actual_delay:.2f}秒待機中...")
            await asyncio.sleep(actual_delay)
    if last_exception:
        raise last_exception

# ==============================================================================
# SECTION 4: モジュール (Modules)
# ==============================================================================

# --- Module 1: DocLoader ---
class DocLoader:
    @staticmethod
    async def load_document(file_path: Path) -> Document:
        suffix = file_path.suffix.lower()
        content = ""
        try:
            if suffix == '.pdf':
                content = await DocLoader._load_pdf(file_path)
            elif suffix in ['.docx', '.doc']:
                content = await DocLoader._load_word(file_path)
            elif suffix in ['.html', '.htm']:
                content = await DocLoader._load_html(file_path)
            elif suffix == '.md':
                content = await DocLoader._load_markdown(file_path)
            else: # .txt
                content = await DocLoader._load_text(file_path)
            metadata = {"filename": file_path.name, "size": file_path.stat().st_size}
            return Document(path=file_path, content=content, metadata=metadata)
        except Exception as e:
            logger.error(f"文書読み込みエラー {file_path}: {e}")
            return Document(path=file_path, content="", metadata={})

    @staticmethod
    async def _load_pdf(file_path: Path) -> str:
        try:
            elements = partition(filename=str(file_path))
            return "\n\n".join([el.text for el in elements if hasattr(el, 'text')])
        except Exception:
            logger.warning(f"unstructuredでのPDF読込失敗。pypdfで再試行します。")
            try:
                with open(file_path, "rb") as f:
                    reader = PdfReader(f)
                    return "\n\n".join([page.extract_text() or "" for page in reader.pages])
            except Exception as e:
                logger.error(f"PDF読み込み完全失敗: {e}")
                return ""

    @staticmethod
    async def _load_word(file_path: Path) -> str:
        doc = docx.Document(str(file_path))
        return "\n\n".join([para.text for para in doc.paragraphs if para.text])

    @staticmethod
    async def _load_html(file_path: Path) -> str:
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html = await f.read()
        soup = BeautifulSoup(html, 'html.parser')
        [s.decompose() for s in soup(['script', 'style'])]
        return soup.get_text(separator='\n', strip=True)

    @staticmethod
    async def _load_markdown(file_path: Path) -> str:
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            md = await f.read()
        html = markdown.markdown(md)
        return BeautifulSoup(html, 'html.parser').get_text(separator='\n', strip=True)

    @staticmethod
    async def _load_text(file_path: Path) -> str:
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return await f.read()
        
# --- Module 2: TextPreprocessor ---
class TextPreprocessor:
    def preprocess(self, doc: Document) -> Document:
        doc.content = self._normalize_text(doc.content)
        doc.chunks = self._split_text_into_chunks(doc.content)
        return doc

    def _normalize_text(self, text: str) -> str:
        if not text: return ""
        text = unicodedata.normalize('NFKC', text)
        text = re.sub(r'\s+', ' ', text).strip()
        text = re.sub(r'\n\s*\n', '\n', text)
        return text

    def _split_text_into_chunks(self, text: str, chunk_size: int = 4000, overlap: int = 400) -> List[str]:
        if not text:
            return []
        sentences = re.split(r'(?<=[。．！？\n])', text)
        chunks, current_chunk = [], ""
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append(current_chunk)
                overlap_start_index = max(0, len(current_chunk) - overlap)
                overlap_text = current_chunk[overlap_start_index:]
                sentence_boundary = re.search(r'[。．！？\n]', overlap_text)
                current_chunk = overlap_text[sentence_boundary.end():] if sentence_boundary else overlap_text
            current_chunk += sentence
        if current_chunk:
            chunks.append(current_chunk)
        final_chunks = [c.strip() for c in chunks if c.strip()]
        logger.info(f"テキストを {len(final_chunks)} 個のチャンクに分割しました")
        return final_chunks

# --- Module 3: LLMTermExtractor (パス1) ---
TERM_LIST_SCHEMA = {
    "type": "object",
    "properties": {
        "terms": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "headword": {"type": "string"},
                    "synonyms": {"type": "array", "items": {"type": "string"}},
                    "definition": {"type": "string"},
                    "category": {"type": "string"}
                },
                "required": ["headword", "definition"]
            }
        }
    },
    "required": ["terms"]
}

class LLMTermExtractor:
    """(パス1) Few-shotプロンプティングを用いてテキストから専門用語候補を抽出する"""
    def __init__(self, config: Config):
        self.config = config
        self.model = genai.GenerativeModel(config.llm_model)

    async def extract_terms_from_chunk(self, chunk: str, doc_name: str) -> List[Term]:
        prompt = self._build_extraction_prompt(chunk)
        response = None
        try:
            logger.info(f"LLM抽出開始(パス1): {doc_name} の一部（文字数: {len(chunk)}）")
            if len(chunk) > 30000:
                chunk = chunk[:30000]
                logger.warning(f"チャンクが長すぎるため30000文字に切り詰めました")

            # 修正: 有効なHarmCategoryのみを使用
            safety_settings = [
                {"category": HarmCategory.HARM_CATEGORY_HARASSMENT, "threshold": HarmBlockThreshold.BLOCK_NONE},
                {"category": HarmCategory.HARM_CATEGORY_HATE_SPEECH, "threshold": HarmBlockThreshold.BLOCK_NONE},
                {"category": HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, "threshold": HarmBlockThreshold.BLOCK_NONE},
                {"category": HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, "threshold": HarmBlockThreshold.BLOCK_NONE},
            ]

            response = await self.model.generate_content_async(
                [prompt],
                generation_config=GenerationConfig(
                    response_mime_type="application/json",
                    response_schema=TERM_LIST_SCHEMA,
                    temperature=0.1,
                    max_output_tokens=8192,
                ),
                safety_settings=safety_settings
            )
            if not response or not response.text:
                logger.warning(f"LLMから空の応答を受信しました(パス1): {doc_name}")
                return []

            result_json = json.loads(response.text)
            term_list = TermList.model_validate(result_json)
            extracted_terms = []
            for t in term_list.terms:
                if not t.headword or len(t.headword) < 2:
                    continue
                term_obj = Term(
                    headword=t.headword.strip(),
                    synonyms=[s.strip() for s in t.synonyms if s.strip()],
                    definition=t.definition.strip() if t.definition else "",
                    category=t.category.strip() if t.category else None,
                    source_docs={doc_name}
                )
                extracted_terms.append(term_obj)
            logger.info(f"チャンクから {len(extracted_terms)} 件の用語候補を抽出しました(パス1)")
            return extracted_terms
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析エラー(パス1): {e}")
            if response: logger.error(f"応答テキスト: {response.text}")
            return []
        except Exception as e:
            logger.error(f"LLM/Pydantic処理エラー(パス1) ({doc_name}): {type(e).__name__}: {e}")
            if response: logger.error(f"LLM応答: {getattr(response, 'text', 'N/A')}")
            return []

    async def extract_terms_from_chunk_with_retry(self, chunk: str, doc_name: str) -> List[Term]:
        async def _extract():
            return await self.extract_terms_from_chunk(chunk, doc_name)
        try:
            result = await retry_async(_extract, max_attempts=3, initial_delay=2.0)
            return result or []
        except Exception as e:
            logger.error(f"最終的に抽出に失敗しました(パス1): {doc_name}: {e}")
            return []

    def _build_extraction_prompt(self, text: str) -> str:
        return f"""
あなたは、与えられた技術文書を分析し、重要な専門用語とその関連情報を構造化する専門家です。
以下のテキストを読み、ドメイン固有の専門用語、その類義語（表記揺れ含む）、文脈に基づいた簡潔な定義、そしてその用語が属する**上位の概念（カテゴリ）**を抽出してください。

# 指示
- 出力は必ず指定されたJSON形式に従ってください。
- 一般的な単語（例：「問題」「解決」「方法」）は含めないでください。
- 技術的な専門性を持つ用語のみを対象とします。
- **定義**: 文脈に基づいて30〜50字程度で簡潔に記述してください。
- **カテゴリ**: その用語が属するより大きなグループ名を指定してください。例えば、「フロントバンパー」のカテゴリは「外装部品」です。適切なカテゴリがない場合はnullにしてください。
- 各用語は、最も基本的な形（例：「ヘッドライト」）で見出し語としてください。

# 例1
[入力テキスト]
...車両前方の外観を構成する部品としてフロントバンパーがある。このフロントバンパーフェイスは樹脂製で...

[期待される出力]
{{
  "terms": [
    {{
      "headword": "フロントバンパー",
      "synonyms": ["フロントバンパーフェイス"],
      "definition": "車両前方の衝撃を吸収する外装部品。",
      "category": "外装部品"
    }}
  ]
}}

# 例2
[入力テキスト]
...リビルト部品のコア返却をお願いします。コアとは、再生の核となる使用済み部品のことです...

[期待される出力]
{{
  "terms": [
    {{
      "headword": "コア",
      "synonyms": ["コア部品"],
      "definition": "リビルト部品の再生・修理の核となる使用済み部品。",
      "category": "リビルト部品"
    }}
  ]
}}

---
それでは、以下のテキストから専門用語を抽出し、指定されたJSON形式で結果を返してください。

[分析対象テキスト]
{text}
"""

# --- Module 4: SynonymLinker (パス2) ---
class SynonymLinker:
    """(パス2) 抽出された全用語候補をLLMで統合・整理する"""
    def __init__(self, config: Config):
        self.config = config
        self.model = genai.GenerativeModel(config.llm_model)
        self.embeddings_client = GoogleGenerativeAIEmbeddings(
            model=config.embedding_model,
            google_api_key=config.google_api_key
        )

    async def consolidate_terms(self, terms: List[Term]) -> List[Term]:
        if not terms:
            return []
        
        console.print("[cyan]第2パス: LLMによる用語の全体統合を開始...[/cyan]")
        
        # 1. パス1の結果をマージし、LLMへの入力を作成
        #    headwordをキーに重複をある程度排除しつつ、情報を集約
        initial_term_map: Dict[str, Dict] = {}
        for term in terms:
            key = term.headword.lower()
            if key not in initial_term_map:
                initial_term_map[key] = {
                    "headword": term.headword,
                    "synonyms": set(term.synonyms),
                    "definitions": {term.definition} if term.definition else set(),
                    "categories": {term.category} if term.category else set(),
                    "source_docs": set(term.source_docs),
                }
            else:
                existing = initial_term_map[key]
                existing["synonyms"].update(term.synonyms)
                if term.definition: existing["definitions"].add(term.definition)
                if term.category: existing["categories"].add(term.category)
                existing["source_docs"].update(term.source_docs)

        # 2. LLMへの入力用に整形
        candidate_list = []
        for data in initial_term_map.values():
            candidate_list.append({
                "headword": data["headword"],
                "synonyms": list(data["synonyms"]),
                "definition": " / ".join(data["definitions"]), # 複数の定義候補を提示
                "category": list(data["categories"])[0] if data["categories"] else None
            })

        candidate_json = json.dumps({"terms": candidate_list}, ensure_ascii=False, indent=2)

        # 3. LLMによる統合（パス2）の実行
        prompt = self._build_consolidation_prompt(candidate_json)
        final_terms = await self._run_consolidation_pass_with_retry(prompt)

        # 4. 最終用語リストのエンベディングを取得
        words_to_embed = [t.headword for t in final_terms]
        embeddings = await self._get_embeddings_batch(words_to_embed)
        for term, emb in zip(final_terms, embeddings):
            if emb:
                term.embedding = np.array(emb)
        
        logger.info(f"第2パス完了。{len(final_terms)}件の用語に統合しました。")
        return final_terms

    async def _run_consolidation_pass(self, prompt: str) -> List[Term]:
        """パス2のLLM呼び出しを実行する内部メソッド"""
        # 修正: 有効なHarmCategoryのみを使用
        safety_settings = [
            {"category": HarmCategory.HARM_CATEGORY_HARASSMENT, "threshold": HarmBlockThreshold.BLOCK_NONE},
            {"category": HarmCategory.HARM_CATEGORY_HATE_SPEECH, "threshold": HarmBlockThreshold.BLOCK_NONE},
            {"category": HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, "threshold": HarmBlockThreshold.BLOCK_NONE},
            {"category": HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, "threshold": HarmBlockThreshold.BLOCK_NONE},
        ]
        
        response = await self.model.generate_content_async(
            [prompt],
            generation_config=GenerationConfig(
                response_mime_type="application/json",
                response_schema=TERM_LIST_SCHEMA,
                temperature=0.2, # 少しだけ創造性を許容
                max_output_tokens=8192,
            ),
            safety_settings=safety_settings
        )
        if not response or not response.text:
            logger.error("LLMから空の応答を受信しました(パス2)")
            return []
        
        result_json = json.loads(response.text)
        term_list = TermList.model_validate(result_json)
        
        # Termオブジェクトに変換
        consolidated_terms = []
        for t in term_list.terms:
            # ここではsource_docsは不明なので、空のままにしておく
            term_obj = Term(
                headword=t.headword.strip(),
                synonyms=[s.strip() for s in t.synonyms if s.strip()],
                definition=t.definition.strip(),
                category=t.category.strip() if t.category else None
            )
            consolidated_terms.append(term_obj)
        return consolidated_terms
    
    async def _run_consolidation_pass_with_retry(self, prompt: str) -> List[Term]:
        try:
            result = await retry_async(self._run_consolidation_pass, prompt=prompt)
            return result or []
        except Exception as e:
            logger.error(f"最終的に統合に失敗しました(パス2): {e}")
            return []

    def _build_consolidation_prompt(self, terms_json: str) -> str:
        return f"""
あなたは、専門用語辞書を編纂する編集長です。
以下に、複数の文書からバラバラに抽出された専門用語の候補リストがJSON形式で与えられます。
あなたの仕事は、このリストを分析し、重複や表記揺れを整理・統合して、一つのクリーンで一貫性のある辞書データにまとめることです。

# 指示
1.  **類義語のグループ化**: リスト全体を見て、同じ意味を持つ用語（表記揺れ、略語、正式名称などを含む）を一つのグループにまとめてください。
2.  **代表語（headword）の選定**: 各グループの中から、最も一般的で代表的だと思われる単語を一つだけ選び、それを `headword` としてください。
3.  **情報の統合**: グループ内のすべての情報を集約してください。
    - `synonyms`: `headword` 以外の単語はすべて `synonyms` リストに入れてください。
    - `definition`: グループ内の複数の定義候補を参考に、最も包括的で分かりやすい定義を一つ生成してください。
    - `category`: グループ内で最も適切、または共通する `category` を一つ選んでください。
4.  **出力**: 最終的に整理された用語リストを、元のJSONスキーマと同じ形式で出力してください。重複した用語は完全に排除してください。

# 候補リスト (JSON)
{terms_json}

---
それでは、上記のリストを整理・統合し、完成した辞書データをJSON形式で返してください。
"""

    async def _get_embeddings_batch(self, texts: List[str]) -> List[Optional[List[float]]]:
        all_embeddings: List[Optional[List[float]]] = []
        if not texts: return all_embeddings
        for i in range(0, len(texts), self.config.batch_size):
            batch = texts[i:i + self.config.batch_size]
            try:
                logger.info(f"Embedding取得中: バッチ {i//self.config.batch_size + 1}/{(len(texts)-1)//self.config.batch_size + 1}")
                batch_embeddings = await self.embeddings_client.aembed_documents(batch)
                all_embeddings.extend(batch_embeddings)
            except Exception as e:
                logger.error(f"Embedding取得バッチエラー: {e}")
                all_embeddings.extend([None] * len(batch))
        return all_embeddings

# --- Module 5: DictionaryStore ---
class DictionaryStore:
    def __init__(self, config: Config):
        self.config = config
        self.conn = None
        if config.storage_driver == "postgres":
            self._init_postgres()

    def _init_postgres(self):
        try:
            self.conn = psycopg2.connect(self.config.dsn)
            if self.config.use_pgvector:
                register_vector(self.conn)
            self._create_tables()
            logger.info("PostgreSQL接続成功")
        except psycopg2.Error as e:
            logger.error(f"PostgreSQL接続エラー: {e}")
            console.print(f"[bold yellow]警告: PostgreSQL接続エラー。JSONファイルのみに保存します。[/bold yellow]")
            self.conn = None

    def _create_tables(self):
        if not self.conn: return
        with self.conn.cursor() as cur:
            if self.config.use_pgvector:
                cur.execute("CREATE EXTENSION IF NOT EXISTS vector;")
            cur.execute(f"""
            CREATE TABLE IF NOT EXISTS term_dictionary (
                id SERIAL PRIMARY KEY,
                headword TEXT NOT NULL UNIQUE,
                synonyms TEXT[] DEFAULT '{{}}',
                definition TEXT,
                category TEXT,
                confidence FLOAT,
                source_docs TEXT[] DEFAULT '{{}}',
                embedding vector({self.config.embedding_dim}),
                created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
            );""")
            self.conn.commit()

    def save_terms(self, terms: List[Term]):
        if not self.conn:
            logger.warning("データベース接続がないため、ファイル保存のみ行います。")
            return
        upsert_count = 0
        with self.conn.cursor() as cur:
            for term in terms:
                try:
                    embedding_val = term.embedding.tolist() if term.embedding is not None else None
                    cur.execute("""
                    INSERT INTO term_dictionary (headword, synonyms, definition, category, confidence, source_docs, embedding, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP)
                    ON CONFLICT (headword) DO UPDATE SET
                        synonyms = ARRAY(SELECT DISTINCT e FROM unnest(term_dictionary.synonyms || EXCLUDED.synonyms) e),
                        definition = COALESCE(NULLIF(EXCLUDED.definition, ''), term_dictionary.definition),
                        category = COALESCE(EXCLUDED.category, term_dictionary.category),
                        confidence = GREATEST(term_dictionary.confidence, EXCLUDED.confidence),
                        source_docs = ARRAY(SELECT DISTINCT e FROM unnest(term_dictionary.source_docs || EXCLUDED.source_docs) e),
                        embedding = COALESCE(EXCLUDED.embedding, term_dictionary.embedding),
                        updated_at = CURRENT_TIMESTAMP;
                    """, (term.headword, term.synonyms, term.definition, term.category, term.confidence, list(term.source_docs), embedding_val))
                    upsert_count += 1
                except Exception as e:
                    logger.error(f"用語保存エラー {term.headword}: {e}")
                    self.conn.rollback()
            self.conn.commit()
        logger.info(f"{upsert_count}件の用語をデータベースに保存/更新しました")

    def export_to_json(self, terms: List[Term], output_path: str):
        dictionary = [term.to_dict() for term in terms]
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(dictionary, f, ensure_ascii=False, indent=2)
        logger.info(f"辞書を{output_path}にエクスポートしました（{len(dictionary)}件）")

    def close(self):
        if self.conn: self.conn.close()

# ==============================================================================
# SECTION 5: パイプライン (Pipeline)
# ==============================================================================

class ExtractionPipeline:
    def __init__(self, config: Config):
        self.config = config
        self.loader = DocLoader()
        self.preprocessor = TextPreprocessor()
        self.llm_extractor = LLMTermExtractor(config)
        self.synonym_linker = SynonymLinker(config)
        self.store = DictionaryStore(config)

    async def process_document(self, file_path: Path) -> List[Term]:
        logger.info(f"処理開始: {file_path.name}")
        doc = await self.loader.load_document(file_path)
        if not doc.content:
            logger.warning(f"文書が空です: {file_path.name}")
            return []
        doc = self.preprocessor.preprocess(doc)
        tasks = [self.llm_extractor.extract_terms_from_chunk_with_retry(chunk, file_path.name) for chunk in doc.chunks]
        results_from_chunks = await asyncio.gather(*tasks)
        doc_terms = [term for chunk_result in results_from_chunks for term in chunk_result]
        logger.info(f"ファイル '{file_path.name}' から {len(doc_terms)} 件の用語候補を抽出しました。")
        return doc_terms

    async def process_directory(self, input_dir: Path, output_dir: Path):
        file_patterns = ["*.pdf", "*.docx", "*.doc", "*.html", "*.htm", "*.md", "*.txt"]
        all_files = [f for p in file_patterns for f in input_dir.glob(p) if f.is_file()]
        logger.info(f"処理対象ファイル数: {len(all_files)}")
        if not all_files:
            console.print("[yellow]処理対象ファイルが見つかりません。[/yellow]")
            return
        
        console.print("[cyan]第1パス: 各ファイルから用語候補を抽出中...[/cyan]")
        semaphore = asyncio.Semaphore(self.config.parallelism)
        async def process_with_semaphore(file):
            async with semaphore:
                try:
                    return await self.process_document(file)
                except Exception as e:
                    logger.error(f"ファイル処理中のパイプラインエラー {file.name}: {e}", exc_info=True)
                    return []
        results = await asyncio.gather(*[process_with_semaphore(f) for f in all_files])
        all_extracted_terms = [term for res in results for term in res]
        logger.info(f"全ファイルからの総抽出用語数（統合前）: {len(all_extracted_terms)}")
        if not all_extracted_terms:
            console.print("[yellow]抽出された用語がありませんでした。[/yellow]")
            return
        
        final_terms = await self.synonym_linker.consolidate_terms(all_extracted_terms)
        
        # パス2で失われたソース情報を復元
        headword_map = {t.headword: t for t in all_extracted_terms}
        for final_term in final_terms:
            if final_term.headword in headword_map:
                 final_term.source_docs = headword_map[final_term.headword].source_docs

        self.store.save_terms(final_terms)
        output_json_path = output_dir / "dictionary.json"
        self.store.export_to_json(final_terms, str(output_json_path))
        self._print_statistics(final_terms)
        self.store.close()

    def _print_statistics(self, terms: List[Term]):
        if not terms:
            console.print("[yellow]最終的な用語リストは空です。[/yellow]")
            return
        table = Table(title="最終抽出結果統計")
        table.add_column("項目", style="cyan")
        table.add_column("値", style="magenta")
        table.add_row("総ユニーク見出し語数", str(len(terms)))
        table.add_row("類義語を持つ用語数", str(sum(1 for t in terms if t.synonyms)))
        table.add_row("定義を持つ用語数", str(sum(1 for t in terms if t.definition)))
        table.add_row("カテゴリを持つ用語数", str(sum(1 for t in terms if t.category)))
        total_synonyms = sum(len(t.synonyms) for t in terms)
        avg_synonyms = f"{(total_synonyms / len(terms)):.2f}" if terms else "0.00"
        table.add_row("平均類義語数", avg_synonyms)
        console.print(table)

# ==============================================================================
# SECTION 6: API (FastAPI) & CLI (Typer)
# ==============================================================================

api = FastAPI(title="専門用語辞書API", version="1.4.0")

def get_store():
    config = Config("config.yml")
    store = DictionaryStore(config)
    try:
        yield store
    finally:
        store.close()

@api.get("/terms")
async def search_terms_api(
    query: str = Query(..., description="検索クエリ"),
    limit: int = Query(10, description="結果数上限"),
    store: DictionaryStore = Depends(get_store)
):
    if not store.conn:
        raise HTTPException(status_code=503, detail="データベースに接続できません")
    try:
        results = []
        with store.conn.cursor() as cur:
            cur.execute("""
            SELECT headword, synonyms, definition, category, confidence, source_docs
            FROM term_dictionary
            WHERE headword ILIKE %s OR %s = ANY(synonyms)
            ORDER BY updated_at DESC, confidence DESC
            LIMIT %s;
            """, (f"%{query}%", query, limit))
            for row in cur.fetchall():
                results.append({
                    "headword": row[0], "synonyms": row[1], "definition": row[2],
                    "category": row[3], "confidence": row[4], "source_docs": row[5]
                })
        return {"results": results}
    except Exception as e:
        logger.error(f"API検索エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.command()
def run_extract(
    input_dir: str = typer.Option("./input", help="入力ディレクトリ"),
    output_dir: str = typer.Option("./output", help="出力ディレクトリ"),
    config_path: str = typer.Option("config.yml", help="設定ファイルパス")
):
    """(2パスLLM) LLMを使って専門用語抽出を実行"""
    console.print("[bold green]専門用語抽出開始 (2パスLLM・階層化モード)[/bold green]")
    try:
        config_obj = Config(config_path)
    except typer.Exit:
        return
    input_path = Path(input_dir)
    if not input_path.exists():
        console.print(f"[red]入力ディレクトリが存在しません: {input_dir}[/red]")
        raise typer.Exit(1)
    output_path = Path(output_dir)
    pipeline = ExtractionPipeline(config_obj)
    asyncio.run(pipeline.process_directory(input_path, output_path))
    console.print("[bold green]処理完了！[/bold green]")

@app.command()
def serve_api(host: str = "0.0.0.0", port: int = 8000):
    """APIサーバーを起動"""
    console.print(f"[bold green]APIサーバー起動: http://{host}:{port}[/bold green]")
    uvicorn.run("main:api", host=host, port=port, reload=True)

@app.command()
def init_project():
    """プロジェクトを初期化し、設定ファイルを生成"""
    console.print("[bold green]プロジェクト初期化[/bold green]")
    for dir_name in ["input", "output", "logs"]:
        Path(dir_name).mkdir(exist_ok=True)
        console.print(f"✓ {dir_name}/ 作成")
    config_template = """
llm:
  model: gemini-2.0-flash
embedding:
  model: models/text-embedding-004
  dim: 768
  batch_size: 100
storage:
  driver: postgres
  dsn: "postgresql://user:password@localhost:5432/term_db"
  use_pgvector: true
parallelism: 4
"""
    with open("config.yml", 'w', encoding='utf-8') as f:
        f.write(config_template.strip())
    console.print("✓ config.yml を作成/更新しました。")
    env_template = """
GOOGLE_API_KEY=your-api-key-here
DATABASE_URL=postgresql://user:password@localhost:5432/term_db
"""
    if not Path(".env").exists():
        with open(".env", 'w', encoding='utf-8') as f:
            f.write(env_template.strip())
        console.print("✓ .env 作成")
    else:
        console.print("✓ .env は既に存在します")
    console.print("\n[bold green]初期化完了！[/bold green]")
    console.print("次に .env ファイルの GOOGLE_API_KEY を設定してください。")

if __name__ == "__main__":
    app()